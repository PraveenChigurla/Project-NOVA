import os
import sys
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import win32com.client

# Target Document Specifications List
TARGET_DOCUMENTS = [
    r"00_Blueprint\00_Foundation\NOVA-SPEC-001_Project_Foundation_Specification.md",
    r"00_Blueprint\01_Product\NOVA-SPEC-002_Product_Requirements_Specification.md",
    r"00_Blueprint\02_Architecture\NOVA-SPEC-003_System_Architecture_Specification.md",
    r"00_Blueprint\02_Architecture\NOVA-SPEC-010_Communication_Framework_Specification.md",
    r"00_Blueprint\04_AI\NOVA-SPEC-004_AI_Core_Specification.md",
    r"00_Blueprint\04_AI\NOVA-SPEC-009_NOVA_Kernel_Specification.md",
    r"00_Blueprint\05_Capabilities\Desktop\Desktop_Capability_Specification.md",
    r"00_Blueprint\05_Capabilities\Browser\Browser_Capability_Specification.md",
    r"00_Blueprint\05_Capabilities\Vision\Vision_Capability_Specification.md",
    r"00_Blueprint\05_Capabilities\Voice\Voice_Capability_Specification.md",
    r"00_Blueprint\03_Engineering\NOVA-ERR-001_Engineering_Review_Report.md",
    r"00_Blueprint\03_Engineering\NOVA-ENG-002_Engineering_Backlog.md",
    r"00_Blueprint\03_Engineering\NOVA-ARR-002_Backlog_Dependency_Analysis.md",
    r"00_Blueprint\03_Engineering\NOVA-REP-002_Repository_Architecture_Specification.md",
    r"00_Blueprint\03_Engineering\NOVA-TECH-001_Technology_Stack_Specification.md",
    r"00_Blueprint\02_Architecture\NOVA-SPEC-011_Capability_Framework_Specification.md",
    r"00_Blueprint\README.md"
]

# Color Palette Definitions
COLOR_PRIMARY_DARK = RGBColor(27, 54, 93)   # Deep Blue #1B365D
COLOR_SECONDARY = RGBColor(75, 107, 148)   # Steel Blue #4B6B94
COLOR_CHARCOAL = RGBColor(51, 51, 51)      # Charcoal #333333

def set_cell_background(cell, fill_hex):
    """Sets background shading color for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding margins in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def parse_markdown_to_docx(md_path, docx_path):
    """Parses standard markdown layout structures into a styled Word .docx."""
    print(f"Parsing Markdown: {md_path} -> {docx_path}")
    
    if not os.path.exists(md_path):
        print(f"Error: Source Markdown file not found: {md_path}")
        return False
        
    doc = Document()
    
    # Configure document margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_table = False
    table_headers = []
    table_rows = []
    
    in_code_block = False
    code_content = []

    for line in lines:
        stripped = line.strip()

        # Handle Code Block bounds
        if stripped.startswith("```"):
            if in_code_block:
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run("\n".join(code_content))
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = COLOR_SECONDARY
                # Add background shading to block
                pPr = p._p.get_or_add_pPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                pPr.append(shd)
                code_content = []
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_content.append(line.rstrip('\r\n'))
            continue

        # Close active tables if separator is encountered
        if stripped.startswith("|") and not stripped.replace("|", "").replace("-", "").replace(" ", ""):
            # Table divider line (e.g. |---|---|) - ignore it
            continue

        # Handle Table parsing
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if not table_headers:
                table_headers = cells
            else:
                table_rows.append(cells)
            continue
        elif in_table:
            # Terminate and render Table
            if table_headers:
                col_count = len(table_headers)
                table = doc.add_table(rows=1, cols=col_count)
                table.autofit = False
                
                # Render Header row
                hdr_cells = table.rows[0].cells
                for i, text in enumerate(table_headers):
                    hdr_cells[i].text = text
                    set_cell_background(hdr_cells[i], "1B365D")
                    set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
                    # Style text
                    run = hdr_cells[i].paragraphs[0].runs[0]
                    run.font.bold = True
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)
                
                # Render Row structures
                for r_idx, row_cells_data in enumerate(table_rows):
                    row_cells = table.add_row().cells
                    bg_color = "F9FBFD" if r_idx % 2 == 1 else "FFFFFF"
                    for col_idx in range(min(col_count, len(row_cells_data))):
                        row_cells[col_idx].text = row_cells_data[col_idx]
                        set_cell_background(row_cells[col_idx], bg_color)
                        set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
                        # Style text
                        p = row_cells[col_idx].paragraphs[0]
                        p.paragraph_format.space_after = Pt(2)
                        for run in p.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(9.5)
                            run.font.color.rgb = COLOR_CHARCOAL

                # Space after table
                p_spacer = doc.add_paragraph()
                p_spacer.paragraph_format.space_before = Pt(12)
                p_spacer.paragraph_format.space_after = Pt(0)
                
            table_headers = []
            table_rows = []
            in_table = False

        if not stripped:
            continue

        # Handle Headings
        if stripped.startswith("#"):
            match = re.match(r"^(#+)\s+(.*)$", stripped)
            if match:
                level = len(match.group(1))
                title = match.group(2)
                
                p = doc.add_paragraph()
                run = p.add_run(title)
                run.font.bold = True
                
                if level == 1:
                    run.font.size = Pt(22)
                    run.font.color.rgb = COLOR_PRIMARY_DARK
                    p.paragraph_format.space_before = Pt(24)
                    p.paragraph_format.space_after = Pt(12)
                    p.paragraph_format.keep_with_next = True
                    # Set alignment center for Cover headers
                    if md_path.endswith("Specification.md") and len(doc.paragraphs) < 5:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif level == 2:
                    run.font.size = Pt(16)
                    run.font.color.rgb = COLOR_PRIMARY_DARK
                    p.paragraph_format.space_before = Pt(18)
                    p.paragraph_format.space_after = Pt(8)
                    p.paragraph_format.keep_with_next = True
                elif level == 3:
                    run.font.size = Pt(13)
                    run.font.color.rgb = COLOR_SECONDARY
                    p.paragraph_format.space_before = Pt(14)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.keep_with_next = True
                else:
                    run.font.size = Pt(11)
                    run.font.color.rgb = COLOR_CHARCOAL
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(4)
                continue

        # Handle Lists (Bulleted or Numbered)
        is_bullet = stripped.startswith("- ") or stripped.startswith("* ")
        is_number = re.match(r"^\d+\.\s+", stripped)
        
        if is_bullet or is_number:
            style = 'List Bullet' if is_bullet else 'List Number'
            text_content = stripped[2:] if is_bullet else re.sub(r"^\d+\.\s+", "", stripped)
            
            p = doc.add_paragraph(style=style)
            p.paragraph_format.space_after = Pt(3)
            # Parse bold/italic wrappers
            _add_markdown_runs(p, text_content)
            continue

        # Handle Standard Paragraphs
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        _add_markdown_runs(p, stripped)

    # Save compiled .docx
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    return True

def _add_markdown_runs(paragraph, text):
    """Parses simple inline markdown markers (**bold**, *italics*) into Word runs."""
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        run.font.color.rgb = COLOR_CHARCOAL
        
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.font.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.font.italic = True
        else:
            run.text = part

def convert_docx_to_pdf(docx_path, pdf_path):
    """Converts a .docx file into a .pdf using MS Word COM Automation."""
    print(f"Converting PDF: {docx_path} -> {pdf_path}")
    word = None
    doc = None
    try:
        abs_docx = os.path.abspath(docx_path)
        abs_pdf = os.path.abspath(pdf_path)
        
        # Initialize MS Word Application via COM COM
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        doc = word.Documents.Open(abs_docx)
        
        # wdFormatPDF = 17
        doc.SaveAs(abs_pdf, FileFormat=17)
        doc.Close()
        print(f"Successfully exported PDF: {pdf_path}")
        return True
    except Exception as e:
        print(f"Error converting {docx_path} to PDF: {e}")
        return False
    finally:
        if doc is not None:
            try:
                doc.Close()
            except:
                pass
        if word is not None:
            try:
                word.Quit()
            except:
                pass

def main():
    root_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    print(f"Target Project Root: {root_dir}")
    
    compiled_count = 0
    
    for relative_path in TARGET_DOCUMENTS:
        md_file = os.path.join(root_dir, relative_path)
        base_name = os.path.splitext(md_file)[0]
        docx_file = base_name + ".docx"
        pdf_file = base_name + ".pdf"
        
        print("\n" + "="*60)
        # 1. Compile Markdown to Word
        success = parse_markdown_to_docx(md_file, docx_file)
        if success:
            # 2. Convert Word to PDF using local Word COM
            pdf_success = convert_docx_to_pdf(docx_file, pdf_file)
            if pdf_success:
                compiled_count += 1
                
    print("\n" + "="*60)
    print(f"Documentation Refactoring compilation complete. {compiled_count} of {len(TARGET_DOCUMENTS)} files processed successfully.")

if __name__ == "__main__":
    main()
